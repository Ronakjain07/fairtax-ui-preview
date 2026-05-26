"""
Universal File Handler Service

Converts multiple file formats (PDF, DOCX, images) to standardized image format
for Vision model processing. Handles low-quality documents with preprocessing.

Supported formats:
- PDF (multi-page)
- DOCX (extracts embedded images and converts text)
- Images: JPEG, PNG, BMP, TIFF, GIF, WebP
- Scanned documents (with enhancement for low-quality)
"""

import io
import logging
from PIL import Image, ImageEnhance, ImageFilter
import mimetypes

logger = logging.getLogger(__name__)

# Supported MIME types
SUPPORTED_MIMES = {
    'application/pdf': 'pdf',
    'image/jpeg': 'image',
    'image/png': 'image',
    'image/bmp': 'image',
    'image/tiff': 'image',
    'image/gif': 'image',
    'image/webp': 'image',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
}


def detect_file_type(file_bytes, mime_type):
    """
    Detect file type from mime_type and file signature.

    Args:
        file_bytes: Raw file bytes
        mime_type: MIME type string

    Returns:
        str: File type ('pdf', 'image', 'docx')
    """
    mime_type = (mime_type or '').lower()

    # Check MIME type first
    if mime_type in SUPPORTED_MIMES:
        return SUPPORTED_MIMES[mime_type]

    # Check magic bytes (file signature)
    if len(file_bytes) >= 4:
        header = file_bytes[:4]

        # PDF: %PDF
        if header.startswith(b'%PDF'):
            return 'pdf'

        # DOCX: PK (ZIP header)
        if header.startswith(b'PK'):
            return 'docx'

        # PNG: PNG signature
        if file_bytes[:8] == b'\x89PNG\r\n\x1a\n':
            return 'image'

        # JPEG: FF D8 FF
        if header[:3] == b'\xff\xd8\xff':
            return 'image'

        # BMP: BM
        if header[:2] == b'BM':
            return 'image'

        # GIF: GIF87a or GIF89a
        if header[:3] in (b'GIF'):
            return 'image'

        # TIFF: II*\0 or MM\0*
        if header[:2] in (b'II', b'MM'):
            return 'image'

    # Fallback to MIME type detection
    if 'pdf' in mime_type:
        return 'pdf'
    elif 'docx' in mime_type or 'word' in mime_type:
        return 'docx'
    elif 'image' in mime_type:
        return 'image'

    raise ValueError(f"Unsupported file type: {mime_type}")


def enhance_image_quality(img, is_scanned=False):
    """
    Enhance image quality for better OCR/Vision extraction.

    Args:
        img: PIL Image object
        is_scanned: bool, if True apply stronger enhancements for scanned docs

    Returns:
        PIL Image object (enhanced)
    """
    try:
        # Ensure RGB mode
        if img.mode in ('RGBA', 'LA', 'P'):
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = rgb_img

        # Resize if too small (improves quality)
        if img.width < 800 or img.height < 600:
            scale = max(800 / img.width, 600 / img.height)
            new_size = (int(img.width * scale), int(img.height * scale))
            img = img.resize(new_size, Image.Resampling.LANCZOS)

        if is_scanned:
            # For scanned documents: stronger enhancement

            # 1. Sharpen
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.5)

            # 2. Increase contrast
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.3)

            # 3. Increase brightness slightly
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(1.1)

            # 4. Apply slight blur to reduce noise, then sharpen
            img = img.filter(ImageFilter.MedianFilter(size=3))
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.2)
        else:
            # For regular images: mild enhancement

            # 1. Slight sharpening
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.1)

            # 2. Mild contrast boost
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.1)

        return img

    except Exception as e:
        logger.warning(f"Image enhancement failed: {str(e)}, continuing without enhancement")
        return img


def convert_pdf_to_images(file_bytes, dpi=300):
    """
    Convert PDF to list of images with quality enhancement.

    Args:
        file_bytes: Raw PDF bytes
        dpi: Resolution for conversion (300 for clarity)

    Returns:
        List[bytes]: JPEG images
    """
    try:
        # Validate file is actually a PDF
        if not file_bytes or len(file_bytes) < 100:
            raise ValueError(f"Invalid file: too small ({len(file_bytes)} bytes)")

        if not file_bytes.startswith(b'%PDF'):
            raise ValueError(f"Invalid PDF: does not start with %PDF header. First bytes: {file_bytes[:20]}")

        # Try PyMuPDF first (fitz)
        try:
            import fitz  # PyMuPDF
            logger.info(f"[PDF_CONVERTER] Opening PDF with PyMuPDF ({len(file_bytes)} bytes)...")
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            logger.info(f"[PDF_CONVERTER] PDF opened successfully, {len(doc)} pages found")
            images = []

            for page_num, page in enumerate(doc):
                try:
                    # Use higher quality pixmap
                    pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
                    img_bytes = pix.tobytes("jpeg")
                    img = Image.open(io.BytesIO(img_bytes))

                    # Enhance for better extraction
                    img = enhance_image_quality(img, is_scanned=False)

                    # Convert to JPEG
                    img_bytes_io = io.BytesIO()
                    img.save(img_bytes_io, format="JPEG", quality=95)
                    images.append(img_bytes_io.getvalue())

                except Exception as e:
                    logger.warning(f"Failed to process PDF page {page_num + 1}: {str(e)}")
                    continue

            doc.close()

            if not images:
                raise ValueError("No pages extracted from PDF")

            logger.info(f"[PDF_PROCESSOR] Converted {len(images)} pages from PDF at {dpi} DPI")
            return images

        except Exception as fitz_error:
            logger.warning(f"[PDF_CONVERTER] PyMuPDF failed: {str(fitz_error)}, trying pypdfium2...")

            # Fallback to pypdfium2
            import pypdfium2 as pdfium
            logger.info(f"[PDF_CONVERTER] Opening PDF with pypdfium2 ({len(file_bytes)} bytes)...")

            # Load PDF directly from bytes
            pdf = pdfium.PdfDocument(file_bytes)
            logger.info(f"[PDF_CONVERTER] PDF opened successfully, {len(pdf)} pages found")

            images = []
            for page_num, page in enumerate(pdf):
                try:
                    # Render page to image with specified DPI
                    # DPI scale: 72 DPI is the default, so multiply by (target_dpi / 72)
                    scale = dpi / 72.0
                    bitmap = page.render(
                        scale=scale,
                        rotation=0,
                        rev_byteorder=True  # Swap BGR to RGB for better compatibility
                    )

                    # Convert bitmap to PIL Image
                    img = bitmap.to_pil()

                    # Ensure RGB mode
                    if img.mode != 'RGB':
                        img = img.convert('RGB')

                    # Enhance for better extraction
                    img = enhance_image_quality(img, is_scanned=False)

                    # Convert to JPEG
                    img_bytes_io = io.BytesIO()
                    img.save(img_bytes_io, format="JPEG", quality=95)
                    images.append(img_bytes_io.getvalue())

                except Exception as e:
                    logger.warning(f"Failed to process PDF page {page_num + 1}: {str(e)}")
                    continue

            if not images:
                raise ValueError("No pages extracted from PDF")

            logger.info(f"[PDF_PROCESSOR] Converted {len(images)} pages from PDF at {dpi} DPI using pypdfium2")
            return images

    except Exception as e:
        logger.error(f"[PDF_PROCESSOR] Error converting PDF: {str(e)}")
        raise


def convert_docx_to_images(file_bytes):
    """
    Convert DOCX to images.

    For DOCX, we extract embedded images and also attempt
    to render the document as images (if possible).

    Args:
        file_bytes: Raw DOCX bytes

    Returns:
        List[bytes]: JPEG images
    """
    try:
        from docx import Document
        from docx.oxml.xmlchemy import OxmlElement

        # Load DOCX
        doc_stream = io.BytesIO(file_bytes)
        doc = Document(doc_stream)

        images = []

        # Extract images from document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    img = Image.open(io.BytesIO(image_bytes))

                    # Enhance quality
                    img = enhance_image_quality(img, is_scanned=True)

                    # Convert to JPEG
                    img_bytes_io = io.BytesIO()
                    img.save(img_bytes_io, format="JPEG", quality=95)
                    images.append(img_bytes_io.getvalue())

                except Exception as e:
                    logger.warning(f"Failed to extract image from DOCX: {str(e)}")
                    continue

        if not images:
            logger.warning("No images found in DOCX, attempting text extraction")
            # Fallback: create image from text
            from PIL import ImageDraw, ImageFont

            try:
                text_content = '\n'.join([
                    paragraph.text for paragraph in doc.paragraphs
                    if paragraph.text.strip()
                ])

                if text_content:
                    # Create image from text
                    img = Image.new('RGB', (1200, 1600), color='white')
                    draw = ImageDraw.Draw(img)

                    y = 50
                    for line in text_content.split('\n')[:50]:  # Limit lines
                        try:
                            draw.text((50, y), line[:100], fill='black')
                            y += 30
                        except Exception:
                            pass

                    img_bytes_io = io.BytesIO()
                    img.save(img_bytes_io, format="JPEG", quality=95)
                    images.append(img_bytes_io.getvalue())
            except Exception as e:
                logger.warning(f"Text extraction from DOCX also failed: {str(e)}")

        if not images:
            raise ValueError("No extractable content found in DOCX")

        logger.info(f"[DOCX_PROCESSOR] Extracted {len(images)} image(s) from DOCX")
        return images

    except Exception as e:
        logger.error(f"[DOCX_PROCESSOR] Error processing DOCX: {str(e)}")
        raise


def convert_image_to_images(file_bytes, mime_type=None):
    """
    Process single image file (may need enhancement for scanned docs).

    Args:
        file_bytes: Raw image bytes
        mime_type: Image MIME type

    Returns:
        List[bytes]: JPEG image (as list with single element)
    """
    try:
        img = Image.open(io.BytesIO(file_bytes))

        # Detect if likely scanned document
        is_scanned = detect_if_scanned(img)

        # Enhance quality
        img = enhance_image_quality(img, is_scanned=is_scanned)

        # Convert to JPEG
        img_bytes_io = io.BytesIO()
        img.save(img_bytes_io, format="JPEG", quality=95)

        logger.info(f"[IMAGE_PROCESSOR] Processed image (scanned={is_scanned})")
        return [img_bytes_io.getvalue()]

    except Exception as e:
        logger.error(f"[IMAGE_PROCESSOR] Error processing image: {str(e)}")
        raise


def detect_if_scanned(img):
    """
    Detect if image is likely a scanned document.

    Uses heuristics:
    - Low color variance (mostly black/white/gray)
    - Aspect ratio typical of documents
    - Grain/noise patterns

    Args:
        img: PIL Image object

    Returns:
        bool: True if likely scanned
    """
    try:
        # Convert to RGB if needed
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Check aspect ratio (documents are usually tall)
        aspect = img.height / img.width if img.width > 0 else 1
        is_document_aspect = 1.0 < aspect < 3.0

        # Check color variance (scanned docs have low variance)
        import numpy as np
        arr = np.array(img)
        color_variance = np.std(arr)
        low_variance = color_variance < 50

        # Check if mostly text-like (high contrast)
        grayscale = img.convert('L')
        grayscale_arr = np.array(grayscale)
        contrast = np.std(grayscale_arr)
        high_contrast = contrast > 40

        is_scanned = is_document_aspect and low_variance and high_contrast

        logger.debug(f"Scanned detection: aspect_ratio={aspect:.2f}, "
                    f"color_variance={color_variance:.1f}, contrast={contrast:.1f}, "
                    f"result={is_scanned}")

        return is_scanned

    except Exception as e:
        logger.warning(f"Scanned detection failed: {str(e)}, assuming not scanned")
        return False


def process_file(file_bytes, mime_type):
    """
    Main entry point: process any supported file format to images.

    Args:
        file_bytes: Raw file bytes
        mime_type: MIME type string

    Returns:
        List[bytes]: JPEG images (list with one or more elements)

    Raises:
        ValueError: If file type unsupported or processing fails
    """
    try:
        # Validate input
        if not file_bytes:
            raise ValueError("File is empty")

        logger.info(f"[FILE_HANDLER] Processing file: {len(file_bytes)} bytes, mime={mime_type}")

        # Detect file type
        file_type = detect_file_type(file_bytes, mime_type)
        logger.info(f"[FILE_HANDLER] Detected file type: {file_type} (mime={mime_type})")
        logger.info(f"[FILE_HANDLER] File header: {file_bytes[:20]}")

        # Convert to images based on type
        if file_type == 'pdf':
            images = convert_pdf_to_images(file_bytes)
        elif file_type == 'docx':
            images = convert_docx_to_images(file_bytes)
        elif file_type == 'image':
            images = convert_image_to_images(file_bytes, mime_type)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        if not images:
            raise ValueError("No images generated from file")

        logger.info(f"[FILE_HANDLER] Generated {len(images)} image(s)")
        return images

    except Exception as e:
        logger.error(f"[FILE_HANDLER] Error processing file: {str(e)}")
        raise
